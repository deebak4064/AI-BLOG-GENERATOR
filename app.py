"""
AI Blog Generator - Flask Application
Created by: Deebak Kumar (deebak4064)
Repository: https://github.com/deebak4064/AI-BLOG-GENERATOR
License: MIT with Attribution Requirement
Please give credit when using or modifying this code.
"""

import os
import csv
import json
import requests
import markdown as _markdown
try:
    # google-genai client (optional dependency)
    from google import genai
    _HAS_GOOGLE_GENAI = True
except Exception:
    genai = None
    _HAS_GOOGLE_GENAI = False
from datetime import datetime
from flask import Flask, request, render_template, redirect, url_for, flash, session
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from dotenv import load_dotenv
from werkzeug.exceptions import HTTPException
from flask import Response, make_response
import re
import io
import time
from uuid import uuid4
from models import db, User, UserBlog, UserStats

# Force load the .env file (allow override so local .env takes precedence here)
load_dotenv(override=True)

# LLM provider configuration - Using Gemini only
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
MODEL = os.getenv("LLM_MODEL", "gemini-2.5-flash")
BLOGS_PER_PAGE = int(os.getenv("BLOGS_PER_PAGE", "1000"))

app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", os.urandom(24))

# Database configuration
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///blogs_app.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

# Flask-Login configuration
login_manager = LoginManager()
login_manager.init_app(app)
login_manager.login_view = 'login'

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Create database tables
with app.app_context():
    db.create_all()
    
    # Attribution notice
    print("\n" + "="*70)
    print("AI Blog Generator - Created by Deepak Kumar")
    print("Repository: https://github.com/deebak4064/AI-BLOG-GENERATOR")
    print("License: MIT with Attribution Requirement")
    print("If you use or modify this code, please give credit to the original author.")
    print("="*70 + "\n")
    
    # Initialize stats for existing users (migration)
    try:
        users = User.query.all()
        for user in users:
            # Check if stats already exist
            existing_stats = UserStats.query.filter_by(user_id=user.id).first()
            if not existing_stats:
                # Count existing blogs for this user
                blog_count = UserBlog.query.filter_by(user_id=user.id).count()
                # Create stats with current blog count
                stats = UserStats(
                    user_id=user.id,
                    total_blogs_generated=blog_count,
                    total_downloads=0
                )
                db.session.add(stats)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"Error initializing stats: {e}")

def detect_category(title, details=""):
    """Detect category from blog title and details."""
    # Category keywords mapping
    categories = {
        'Tech': ['python', 'javascript', 'coding', 'programming', 'development', 'web', 'cloud', 'ai', 'machine learning', 'api', 'database', 'devops', 'cybersecurity', 'blockchain', 'docker', 'react', 'node', 'software', 'code', 'algorithm', 'data science', 'computer'],
        'Beauty': ['skincare', 'makeup', 'beauty', 'cosmetics', 'hair', 'nail', 'skincare routine', 'anti-aging', 'cruelty-free', 'wellness'],
        'Education': ['learning', 'study', 'education', 'course', 'student', 'teaching', 'school', 'university', 'academic', 'language', 'skill'],
        'Gaming': ['gaming', 'game', 'esports', 'video game', 'console', 'pc game', 'mobile game', 'gaming setup', 'streaming', 'vr'],
        'Health': ['fitness', 'workout', 'health', 'diet', 'nutrition', 'exercise', 'mental health', 'yoga', 'meditation', 'wellness', 'sleep'],
        'Travel': ['travel', 'trip', 'destination', 'hotel', 'vacation', 'tourism', 'adventure', 'backpack', 'tour', 'explore'],
        'Lifestyle': ['lifestyle', 'minimalist', 'living', 'habits', 'daily routine', 'personal growth', 'productivity', 'home', 'organization', 'digital detox'],
        'Business': ['business', 'startup', 'entrepreneurship', 'marketing', 'leadership', 'finance', 'management', 'strategy', 'sales', 'remote work']
    }
    
    # Combine title and details for analysis
    text = (title + " " + details).lower()
    
    # Find which category has the most matches
    max_matches = 0
    detected_category = 'General'
    
    for category, keywords in categories.items():
        matches = sum(1 for keyword in keywords if keyword in text)
        if matches > max_matches:
            max_matches = matches
            detected_category = category
    
    return detected_category

def generate_blog(title, details="", api_key=None):
    """Generate a blog using Gemini API.

    This function uses the Gemini API and expects GEMINI_API_KEY to be configured.
    The function will raise a descriptive ValueError if the API key is missing.
    """
    if not title or not title.strip():
        raise ValueError("Title cannot be empty")

    # Decide whether the requested topic is technical/programming-focused or general.
    # This influences the wording of system prompts so we don't force a programming
    # perspective for non-technical topics (e.g., 'skincare blog').
    tech_keywords = [
        'program', 'programming', 'python', 'java', 'javascript', 'developer', 'software', 'engineer',
        'code', 'coding', 'api', 'machine learning', 'ml', 'ai', 'artificial intelligence', 'react',
        'django', 'flask', 'node', 'rust', 'go', 'c++', 'c#'
    ]
    combined = (title or '') + ' ' + (details or '')
    lc = combined.lower()
    is_technical = any(k in lc for k in tech_keywords)

    # Use Gemini API
    key = GEMINI_API_KEY
    if not key:
        raise ValueError("GEMINI_API_KEY is not configured. Please provide GEMINI_API_KEY in .env")

    # Prefer the official google-genai client when available.
    model = MODEL or "text-bison-001"
    if is_technical:
        prompt_text = f"Write a high-quality programming blog article titled '{title}'. {details}"
    else:
        prompt_text = f"Write a high-quality blog article titled '{title}'. {details}"

    if _HAS_GOOGLE_GENAI:
        # The google-genai client reads the API key from env by default, but allow explicit key.
        try:
            # create client and call models.generate_content (matches user's snippet)
            client = genai.Client() if getattr(genai, 'Client', None) else None
            if client is None:
                # fallback to module-level helper if provided
                response = genai.generate_text(model=model, prompt=prompt_text)
                return getattr(response, 'text', str(response))

            # If client supports passing api_key at construction, prefer env or key
            # Client typically pulls key from environment variable, so ensure it's present.
            # We still allow constructing with key if the client supports it.
            try:
                # some versions accept api_key parameter
                client = genai.Client(api_key=key)
            except TypeError:
                # older/newer versions may not accept api_key; rely on env
                pass

            # Use models.generate_content if available (per snippet). Fallback to generate or generate_text.
            if hasattr(client.models, 'generate_content'):
                resp = client.models.generate_content(model=model, contents=prompt_text)
                return getattr(resp, 'text', str(resp))
            if hasattr(client.models, 'generate'):
                resp = client.models.generate(model=model, prompt=prompt_text)
                # various client versions return different shapes
                return getattr(resp, 'text', getattr(resp, 'output', str(resp)))

            # last-resort: raise if client present but no known method
            raise RuntimeError('google-genai client present but no supported generate method found')
        except Exception as e:
            raise RuntimeError(f"Failed to generate blog via Gemini (google-genai): {str(e)}")

    # If google-genai is not installed, fall back to REST call
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={key}"
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{"parts": [{"text": prompt_text}]}],
        "generationConfig": {
            "temperature": 0.7,
            "maxOutputTokens": 1024
        }
    }
    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        j = resp.json()
        if isinstance(j, dict):
            if "candidates" in j and j["candidates"]:
                content = j["candidates"][0].get("content", {})
                if "parts" in content and content["parts"]:
                    return content["parts"][0].get("text", "")
                return content.get("text", "")
        return json.dumps(j)
    except requests.RequestException as e:
        raise RuntimeError(f"Failed to generate blog via Gemini REST: {str(e)}")


# Fetch the chat assistant API key from the .env file
CHAT_ASSISTANT_API_KEY = os.getenv("CHAT_ASSISTANT_API_KEY")

def llm_chat(message: str, api_key: str = CHAT_ASSISTANT_API_KEY, max_tokens: int = 1000) -> str:
    """Generate a short assistant reply to a user message using Gemini API.

    This is a lightweight wrapper intended for conversational/assistant responses and
    should NOT be used for full blog generation (use generate_blog for that).
    """
    key = (api_key or CHAT_ASSISTANT_API_KEY or GEMINI_API_KEY)
    if key:
        key = key.strip()
    if not key:
        raise ValueError("GEMINI_API_KEY is not configured for chat responses")
    
    # Detect if user is asking for multiple options (looking for numbers like "3", "5", etc.)
    import re
    asks_for_number = bool(re.search(r'\b([2-9]|\d{2,})\b', message.lower()))
    
    if asks_for_number:
        # User asked for multiple options - format as numbered list
        system_prompt = """You are an expert writing and content assistant specializing in creating compelling titles, headlines, and alternative phrasings.

IMPORTANT INSTRUCTIONS:
- When asked for alternatives, provide COMPLETE, FULL-LENGTH options
- Do NOT truncate or shorten options - provide the complete version of each
- Number each option clearly (1., 2., 3., etc.)
- Each option should be a complete, ready-to-use alternative
- Provide diverse variations that differ in tone, structure, and emphasis
- Think deeply about what would resonate with the target audience
- Make each option meaningfully different and unique from others
- Keep options professional and well-written

FORMAT YOUR RESPONSE:
1. [Complete first option here - full text]
2. [Complete second option here - full text]
3. [Complete third option here - full text]

Always provide COMPLETE OPTIONS, never truncate them."""
    else:
        # User asked for single alternative - provide just one response
        system_prompt = """You are an expert writing and content assistant specializing in creating compelling titles, headlines, and alternative phrasings.

IMPORTANT INSTRUCTIONS:
- Provide a SINGLE, high-quality alternative or improvement
- Do NOT provide multiple options unless explicitly asked
- Provide the complete, full-length response
- Keep the tone professional and natural
- Make the alternative meaningfully different from the original

Provide ONLY ONE response, no numbering, no "Option 1:", just the improved text directly."""
    
    enhanced_message = f"{system_prompt}\n\nUser request:\n{message}"
    
    # Try google-genai client if available
    if _HAS_GOOGLE_GENAI:
        try:
            client = genai.Client(api_key=key)
            if hasattr(client.models, 'generate_content'):
                resp = client.models.generate_content(model=MODEL, contents=enhanced_message)
                if hasattr(resp, 'text') and resp.text:
                    return resp.text
        except Exception as e:
            # Fall through to REST if google-genai fails
            pass

    # REST fallback: Use generateContent endpoint for chat
    prompt_text = enhanced_message
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL or 'gemini-2.5-flash'}:generateContent?key={key}"
    headers = {"Content-Type": "application/json"}
    payload = {
        "contents": [{"parts": [{"text": prompt_text}]}],
        "generationConfig": {
            "temperature": 0.7,
            "maxOutputTokens": max_tokens
        }
    }
    try:
        r = requests.post(url, headers=headers, json=payload, timeout=20)
        r.raise_for_status()
        j = r.json()
        if isinstance(j, dict):
            if 'candidates' in j and j['candidates']:
                content = j['candidates'][0].get('content', {})
                if 'parts' in content and content['parts']:
                    text = content['parts'][0].get('text', '')
                    if text:
                        return text
        return ''
    except requests.RequestException as e:
        raise RuntimeError(f"Gemini REST chat error: {e}")

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        try:
            inputs = request.form.get("blog_inputs", "").strip()
            if not inputs:
                flash("Please enter at least one blog title", "error")
                return redirect(url_for("index"))

            # default to inline display (we provide downloads after generation)
            save_format = request.form.get("save_format", "inline")
            blogs = []
            
            # Expect titles in lines, option for details as "Title | Details"
            for line in inputs.splitlines():
                if not line.strip():
                    continue
                components = line.split('|', 1)
                title = components[0].strip()
                details = components[1].strip() if len(components) > 1 else ""

                try:
                    body = generate_blog(title, details)
                    # Convert plain/markdown text into safe HTML for rendering
                    try:
                        body_html = _markdown.markdown(body, extensions=["fenced_code", "tables"])
                    except Exception:
                        # Fallback: simple newline -> paragraphs
                        paragraphs = [f"<p>{p.strip()}</p>" for p in body.split('\n\n') if p.strip()]
                        body_html = "\n".join(paragraphs)

                    blog = {
                        "title": title,
                        "details": details,
                        "body": body,
                        "body_html": body_html,
                        "date": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    }
                    # filename base for downloads
                    try:
                        blog["filename_base"] = _slugify(title)
                    except Exception:
                        blog["filename_base"] = _slugify(str(title))
                    blogs.append(blog)
                except Exception as e:
                    flash(f"Error generating blog '{title}': {str(e)}", "error")
                    continue

            if not blogs:
                flash("No blogs were generated successfully", "error")
                return redirect(url_for("index"))

            # Save as requested
            try:
                if save_format == "csv":
                    with open("blogs.csv", "w", newline="", encoding="utf-8") as f:
                        w = csv.DictWriter(f, fieldnames=["title", "details", "body", "date"])
                        w.writeheader()
                        w.writerows(blogs)
                    flash("Blogs saved successfully as CSV", "success")
                    return redirect(url_for("show_blog", format="csv"))
                elif save_format == "json":
                    with open("blogs.json", "w", encoding="utf-8") as f:
                        json.dump(blogs, f, ensure_ascii=False, indent=2)
                    flash("Blogs saved successfully as JSON", "success")
                    return redirect(url_for("show_blog", format="json"))
                else:  # Just show on page
                    # Persist generated blogs server-side to avoid exceeding cookie size.
                    try:
                        os.makedirs('data', exist_ok=True)
                        fname = f"last_blogs_{int(time.time())}_{uuid4().hex}.json"
                        fpath = os.path.join('data', fname)
                        with open(fpath, 'w', encoding='utf-8') as fh:
                            json.dump(blogs, fh, ensure_ascii=False, indent=2)
                        session['last_blogs_file'] = fpath
                    except Exception:
                        # fallback to session (may fail for large content)
                        try:
                            session['last_blogs'] = json.dumps(blogs)
                        except Exception:
                            pass

                    # annotate inline-generated blogs with _idx for downloads
                    for i, b in enumerate(blogs):
                        if isinstance(b, dict):
                            b['_idx'] = i
                    
                            # persist generated blog to DB for authenticated users
                            if current_user and getattr(current_user, 'is_authenticated', False):
                                for b_save in blogs:
                                    try:
                                        # Detect category from title and details
                                        category = detect_category(b_save.get('title', ''), b_save.get('details', ''))
                                        ub = UserBlog(
                                            user_id=current_user.id,
                                            title=b_save.get('title') or '',
                                            details=b_save.get('details') or '',
                                            body=b_save.get('body') or '',
                                            body_html=b_save.get('body_html') or '',
                                            filename_base=b_save.get('filename_base') or None,
                                            category=category
                                        )
                                        db.session.add(ub)
                                    except Exception:
                                        pass
                                try:
                                    db.session.commit()
                                except Exception as e:
                                    db.session.rollback()
                                    print(f"Error saving blogs: {e}")
                    return render_template("blog_list.html", blogs=blogs, total_blogs=len(blogs))
            except IOError as e:
                flash(f"Error saving blogs: {str(e)}", "error")
                return render_template("blog_list.html", blogs=blogs, total_blogs=len(blogs))

        except Exception as e:
            print(f"Blog generation error: {type(e).__name__}: {str(e)}")
            import traceback
            traceback.print_exc()
            flash(f"An unexpected error occurred: {str(e)}", "error")
            return redirect(url_for("index"))

    return render_template("index.html")

@app.route("/blog/<format>")
def show_blog(format):
    try:
        page = request.args.get('page', 1, type=int)
        blogs = []
        
        try:
            if format == "csv":
                with open("blogs.csv", encoding="utf-8") as f:
                    r = csv.DictReader(f)
                    blogs = list(r)
            elif format == "json":
                with open("blogs.json", encoding="utf-8") as f:
                    blogs = json.load(f)
        except FileNotFoundError:
            flash(f"No {format.upper()} file found", "error")
            return redirect(url_for("index"))
        except Exception as e:
            flash(f"Error reading {format.upper()} file: {str(e)}", "error")
            return redirect(url_for("index"))

        # If the repository includes an example/sample JSON (e.g. a single
        # example post such as "how to use docker"), don't surface that as
        # the application's blog list by default. Treat known sample titles
        # as absent so users see an empty state until they generate real posts.
        sample_titles = {"how to use docker"}
        if isinstance(blogs, list) and blogs:
            try:
                # normalize titles and check whether every loaded entry is a sample
                all_sample = all(
                    isinstance(b, dict) and b.get("title", "").strip().lower() in sample_titles
                    for b in blogs
                )
                if all_sample:
                    blogs = []
            except Exception:
                pass

        # Sort blogs by date if available
        if blogs and isinstance(blogs, list) and "date" in blogs[0]:
            blogs.sort(key=lambda x: x["date"], reverse=True)

        # if there are no blogs after filtering, render empty state
        if not blogs:
            return render_template("blog_list.html", blogs=[])

        # annotate global index for each blog so download links can reference the original index
        for i, b in enumerate(blogs):
            try:
                b["_idx"] = i
            except Exception:
                # if b is not a dict (unlikely), skip
                pass
        # ensure filename_base exists for each blog
        for b in blogs:
            if isinstance(b, dict) and "filename_base" not in b:
                try:
                    b["filename_base"] = _slugify(b.get("title", "blog"))
                except Exception:
                    b["filename_base"] = _slugify(str(b.get("title", "blog")))

        # Implement pagination
        total_pages = (len(blogs) + BLOGS_PER_PAGE - 1) // BLOGS_PER_PAGE
        page = max(1, min(page, total_pages))
        start_idx = (page - 1) * BLOGS_PER_PAGE
        end_idx = start_idx + BLOGS_PER_PAGE
        current_blogs = blogs[start_idx:end_idx]

        return render_template(
            "blog_list.html",
            blogs=current_blogs,
            format=format,
            page=page,
            total_pages=total_pages,
            total_blogs=len(blogs)
        )
    except Exception as e:
        flash(f"An unexpected error occurred: {str(e)}", "error")
        return redirect(url_for("index"))


@app.route("/blog")
def my_blogs():
    """Show generated blog history for the logged-in user (title list only)."""
    if not current_user or not getattr(current_user, 'is_authenticated', False):
        flash("Please log in to view your blog history", "error")
        return redirect(url_for('login'))

    try:
        page = request.args.get('page', 1, type=int)
        q = UserBlog.query.filter_by(user_id=current_user.id).order_by(UserBlog.created_at.desc())
        total = q.count()
        total_pages = (total + BLOGS_PER_PAGE - 1) // BLOGS_PER_PAGE if total else 0

        if total_pages and page > total_pages:
            page = total_pages
        if page < 1:
            page = 1

        items = q.offset((page - 1) * BLOGS_PER_PAGE).limit(BLOGS_PER_PAGE).all()
        blogs = [{'id': b.id, 'title': b.title, 'category': b.category or 'General', 'created_at': b.created_at.strftime('%Y-%m-%d %H:%M:%S') if b.created_at else None} for b in items]

        return render_template(
            "blog_history.html",
            blogs=blogs,
            page=page,
            total_pages=total_pages,
            total_blogs=total
        )
    except Exception as e:
        flash(f"Unable to load your blogs: {e}", "error")
        return redirect(url_for('index'))


@app.route("/blog/<int:blog_id>")
def view_blog(blog_id):
    """View a single blog entry by ID."""
    if not current_user or not getattr(current_user, 'is_authenticated', False):
        flash("Please log in to view your blogs", "error")
        return redirect(url_for('login'))

    try:
        blog = UserBlog.query.filter_by(id=blog_id, user_id=current_user.id).first()
        if not blog:
            flash("Blog not found", "error")
            return redirect(url_for('my_blogs'))

        b = blog.to_dict()
        b['_idx'] = 0
        return render_template(
            "blog_list.html",
            blogs=[b],
            is_single_view=True
        )
    except Exception as e:
        flash(f"Unable to load blog: {e}", "error")
        return redirect(url_for('my_blogs'))

@app.route("/api/delete-blog/<int:blog_id>", methods=["DELETE"])
def delete_blog(blog_id):
    """Delete a single blog by ID."""
    if not current_user or not getattr(current_user, 'is_authenticated', False):
        return {'error': 'Unauthorized'}, 401

    try:
        blog = UserBlog.query.filter_by(id=blog_id, user_id=current_user.id).first()
        if not blog:
            return {'error': 'Blog not found'}, 404
        
        db.session.delete(blog)
        db.session.commit()
        return {'success': True, 'message': 'Blog deleted successfully'}, 200
    except Exception as e:
        db.session.rollback()
        return {'error': str(e)}, 500

@app.route("/api/clear-all-blogs", methods=["DELETE"])
def clear_all_blogs():
    """Delete all blogs for the current user."""
    if not current_user or not getattr(current_user, 'is_authenticated', False):
        return {'error': 'Unauthorized'}, 401

    try:
        UserBlog.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        return {'success': True, 'message': 'All blogs cleared successfully'}, 200
    except Exception as e:
        db.session.rollback()
        return {'error': str(e)}, 500

@app.errorhandler(Exception)
def handle_error(e):
    error_msg = "An unexpected error occurred"
    if isinstance(e, HTTPException):
        error_msg = e.description
    elif str(e):
        error_msg = str(e)
    flash(error_msg, "error")
    return redirect(url_for("index"))


def _slugify(text: str) -> str:
    s = text.lower()
    s = re.sub(r"[^a-z0-9\-\_]+", "_", s)
    s = re.sub(r"_+", "_", s).strip("_")
    return s[:120] or "blog"


def _sanitize_blog_content(blog: dict) -> dict:
    """Sanitize blog body/body_html to remove accidental instruction prompts.

    Some LLM uses or mis-invocations can cause the saved body to begin with the
    original instruction (e.g. "Write a high-quality blog article titled ...").
    This helper strips a leading paragraph that looks like an instruction so
    downloads contain only the article title and content.
    """
    def _strip_text_prefix(text: str) -> str:
        if not text:
            return text
        parts = text.split('\n\n')
        if parts:
            first = parts[0].strip().lower()
            triggers = ('write a ', 'write an ', 'you are a ', 'write a high-quality', 'write an engaging')
            if any(t in first for t in triggers) and len(parts) > 1:
                return '\n\n'.join(parts[1:]).lstrip()
        return text

    def _strip_html_prefix(html: str) -> str:
        try:
            from bs4 import BeautifulSoup
        except Exception:
            # if bs4 not available, fallback to simple heuristic
            return _strip_text_prefix(re.sub('<[^<]+?>', '', html))
        soup = BeautifulSoup(html, 'html.parser')
        # find first non-empty child
        first = None
        for el in soup.contents:
            if getattr(el, 'name', None) is None:
                text = str(el).strip()
                if text:
                    first = text
                    break
            else:
                txt = el.get_text().strip()
                if txt:
                    first = txt
                    break
        if first:
            first_l = first.lower()
            triggers = ('write a ', 'write an ', 'you are a ', 'write a high-quality', 'write an engaging')
            if any(t in first_l for t in triggers):
                # remove that first element
                # create a new soup without the first top-level element
                new_contents = []
                removed = False
                for el in soup.contents:
                    if not removed:
                        txt = el.get_text().strip() if getattr(el, 'name', None) else str(el).strip()
                        if txt:
                            removed = True
                            continue
                        else:
                            continue
                    new_contents.append(el)
                new_soup = BeautifulSoup('', 'html.parser')
                for el in new_contents:
                    new_soup.append(el)
                return str(new_soup)
        return html

    # Work on a shallow copy to avoid mutating input unexpectedly
    out = dict(blog)
    try:
        out['body'] = _strip_text_prefix(out.get('body', ''))
    except Exception:
        pass
    try:
        out['body_html'] = _strip_html_prefix(out.get('body_html', ''))
    except Exception:
        pass
    return out


@app.route('/download/<fmt>/<int:idx>')
def download_blog(fmt, idx):
    """Download a blog in specified format.
    
    Supports both:
    1. Session/inline blogs (idx = array index)
    2. Database blogs (idx = UserBlog.id when user is authenticated)
    """
    blog = None
    
    # If user is authenticated, try to load from database first
    if current_user and getattr(current_user, 'is_authenticated', False):
        try:
            blog_obj = UserBlog.query.filter_by(id=idx, user_id=current_user.id).first()
            if blog_obj:
                blog = blog_obj.to_dict()
        except Exception:
            pass
    
    # Fallback to session-based blogs (for inline generation)
    if not blog:
        blogs = []
        fpath = session.get('last_blogs_file')
        if fpath and os.path.exists(fpath):
            try:
                with open(fpath, 'r', encoding='utf-8') as fh:
                    blogs = json.load(fh)
            except Exception:
                blogs = []
        else:
            # fallback to session data
            try:
                blogs = json.loads(session.get('last_blogs', '[]'))
            except Exception:
                blogs = []
        
        try:
            blog = blogs[idx]
        except Exception:
            pass
    
    if not blog:
        flash("Requested blog not found", "error")
        return redirect(url_for('index'))

    # Sanitize to remove accidental instruction prefixes saved in the body
    try:
        blog = _sanitize_blog_content(blog)
    except Exception:
        pass

    # Use input title as filename base only, don't include it in content
    title = blog.get('title', 'blog')
    filename_base = _slugify(title)

    if fmt == 'html':
        # Download only AI-generated content (body_html), not the input title
        content = f"<!doctype html><html><head><meta charset=\"utf-8\"></head><body>{blog.get('body_html','')}</body></html>"
        mimetype = 'text/html'
        filename = f"{filename_base}.html"
    elif fmt == 'pdf':
        # generate PDF bytes - use only generated content
        text = blog.get('body', '')
        html_content = blog.get('body_html') or None
        # Don't include input title in PDF
        pdf_bytes = _render_text_to_pdf_bytes('', text, html=html_content)
        resp = make_response(pdf_bytes)
        # respect custom download_name if provided
        download_name = request.args.get('download_name')
        out_name = filename_base + '.pdf'
        if download_name:
            # sanitize: strip path separators
            safe = os.path.basename(download_name)
            # if extension missing, add .pdf
            if not os.path.splitext(safe)[1]:
                safe = safe + '.pdf'
            out_name = safe
        resp.headers.set('Content-Type', 'application/pdf')
        resp.headers.set('Content-Disposition', f'attachment; filename="{out_name}"')
        return resp
    elif fmt == 'docx':
        # Use only generated content
        docx_bytes = _render_text_to_docx_bytes('', blog.get('body', ''))
        resp = make_response(docx_bytes)
        download_name = request.args.get('download_name')
        out_name = filename_base + '.docx'
        if download_name:
            safe = os.path.basename(download_name)
            if not os.path.splitext(safe)[1]:
                safe = safe + '.docx'
            out_name = safe
        resp.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        resp.headers.set('Content-Disposition', f'attachment; filename="{out_name}"')
        return resp
    elif fmt == 'md' or fmt == 'markdown':
        content = blog.get('body', '')
        mimetype = 'text/markdown'
        filename = f"{filename_base}.md"
    elif fmt == 'txt':
        content = blog.get('body', '')
        mimetype = 'text/plain'
        filename = f"{filename_base}.txt"
    elif fmt == 'json':
        content = json.dumps(blog, ensure_ascii=False, indent=2)
        mimetype = 'application/json'
        filename = f"{filename_base}.json"
    elif fmt == 'csv':
        # single-row CSV
        out = 'title,details,body,date\n'
        def esc(s):
            return '"' + (s or '').replace('"', '""') + '"'
        out += f"{esc(blog.get('title'))},{esc(blog.get('details'))},{esc(blog.get('body'))},{esc(blog.get('date'))}\n"
        content = out
        mimetype = 'text/csv'
        filename = f"{filename_base}.csv"
    else:
        flash('Unsupported download format', 'error')
        return redirect(url_for('index'))

    resp = make_response(content)
    # allow user-specified filename via query param for non-binary formats too
    download_name = request.args.get('download_name')
    out_name = filename
    if download_name:
        safe = os.path.basename(download_name)
        # ensure extension matches desired format
        ext = os.path.splitext(filename)[1]
        if not os.path.splitext(safe)[1]:
            safe = safe + ext
        out_name = safe

    resp.headers.set('Content-Type', mimetype + '; charset=utf-8')
    resp.headers.set('Content-Disposition', f'attachment; filename="{out_name}"')
    return resp


def _render_text_to_pdf_bytes(title: str, text: str, html: str = None) -> bytes:
    """Render PDF bytes.

    Prefer WeasyPrint (HTML -> PDF) for highest fidelity if available. If WeasyPrint
    is not installed, fall back to the existing reportlab-based renderer.

    Accepts either raw markdown/text in `text` or pre-rendered HTML in `html`.
    """
    # First try WeasyPrint (best fidelity from HTML)
    try:
        from weasyprint import HTML, CSS
        have_weasy = True
    except Exception:
        have_weasy = False

    if have_weasy:
        # Build HTML content: prefer provided html, otherwise convert markdown
        try:
            if not html:
                html = _markdown.markdown(text or '', extensions=["fenced_code", "tables", "nl2br"])
            # Basic stylesheet for nicer output
            css_text = '''
                @page { size: A4; margin: 1in; }
                body { font-family: Arial, Helvetica, sans-serif; font-size: 12pt; color: #222; }
                h1, h2, h3 { color: #111; }
                pre { background: #f5f5f5; padding: 8px; border-radius:4px; }
                code { font-family: "Courier New", monospace; }
                table { border-collapse: collapse; }
                table, th, td { border: 1px solid #ccc; }
                th, td { padding: 6px; }
            '''
            # Ensure title included at top
            full_html = f"<html><head><meta charset=\"utf-8\"><title>{title}</title></head><body><h1>{title}</h1>{html}</body></html>"
            pdf_bytes = HTML(string=full_html).write_pdf(stylesheets=[CSS(string=css_text)])
            return pdf_bytes
        except Exception as e:
            # Fall back to reportlab if weasy fails at runtime
            # continue on to the fallback below
            pass

    # Fallback: reportlab-based renderer (keeps previous behavior)
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except Exception:
        raise RuntimeError('reportlab is required to generate PDFs when WeasyPrint is not available: pip install reportlab or weasyprint')

    # We'll render markdown-aware content by converting to HTML and walking the nodes
    from bs4 import BeautifulSoup
    html_content = html or _markdown.markdown(text or '', extensions=["fenced_code", "tables", "nl2br"])
    soup = BeautifulSoup(html_content, "html.parser")

    buf = io.BytesIO()
    c = canvas.Canvas(buf, pagesize=letter)
    width, height = letter
    left_margin = 72
    right_margin = 72
    top_margin = 72
    bottom_margin = 72
    usable_width = width - left_margin - right_margin
    y = height - top_margin

    def ensure_space(h):
        nonlocal y
        if y - h < bottom_margin:
            c.showPage()
            # reset y on new page
            y = height - top_margin

    def draw_wrapped(text, font_name='Helvetica', font_size=10, indent=0, leading=None):
        """Draw `text` wrapped to the usable width using the given font and size.

        This measures string widths and builds lines that fit the available width.
        """
        nonlocal y
        if leading is None:
            leading = int(font_size * 1.2)

        # set font
        try:
            c.setFont(font_name, font_size)
        except Exception:
            c.setFont('Helvetica', font_size)

        words = text.split()
        if not words:
            return

        cur_line = ''
        for word in words:
            test_line = (cur_line + ' ' + word).strip() if cur_line else word
            width_px = pdfmetrics.stringWidth(test_line, font_name, font_size)
            if width_px + indent > usable_width:
                # emit cur_line
                ensure_space(leading)
                c.drawString(left_margin + indent, y, cur_line)
                y -= leading
                cur_line = word
            else:
                cur_line = test_line

        # emit last line
        if cur_line:
            ensure_space(leading)
            c.drawString(left_margin + indent, y, cur_line)
            y -= leading

    # Title
    title_font = 'Helvetica-Bold'
    title_size = 16
    ensure_space(title_size + 6)
    c.setFont(title_font, title_size)
    c.drawString(left_margin, y, title)
    y -= title_size + 12

    # Walk simple tags and use measured wrapping
    for node in soup.children:
        name = getattr(node, 'name', None)
        if name in ('h1', 'h2', 'h3'):
            size = 14 if name == 'h2' else 12 if name == 'h3' else 16
            c.setFont('Helvetica-Bold', size)
            draw_wrapped(node.get_text(), font_name='Helvetica-Bold', font_size=size, indent=0, leading=int(size * 1.3))
            y -= 6
            c.setFont('Helvetica', 10)
        elif name == 'pre' or name == 'code':
            # code block: use monospace and preserve indentation
            code_text = node.get_text()
            c.setFont('Courier', 9)
            # handle lines individually, wrapping long lines by splitting
            for line in code_text.split('\n'):
                # split long continuous lines into chunks that fit
                if not line:
                    ensure_space(12)
                    y -= 12
                    continue
                # try to wrap by words; if still too long, hard-split
                parts = []
                cur = ''
                for ch in line:
                    test = cur + ch
                    if pdfmetrics.stringWidth(test, 'Courier', 9) > usable_width - 24:
                        parts.append(cur)
                        cur = ch
                    else:
                        cur = test
                if cur:
                    parts.append(cur)
                for p in parts:
                    ensure_space(12)
                    c.drawString(left_margin + 12, y, p)
                    y -= 12
            y -= 8
            c.setFont('Helvetica', 10)
        elif name in ('ul', 'ol'):
            for li in node.find_all('li', recursive=False):
                bullet = '\u2022' if name == 'ul' else f"{li.parent.index(li)+1}."
                text = f"{bullet} {li.get_text()}"
                draw_wrapped(text, font_name='Helvetica', font_size=10, indent=8)
            y -= 6
        else:
            # paragraphs or other text
            text_val = node.get_text().strip()
            if not text_val:
                continue
            draw_wrapped(text_val, font_name='Helvetica', font_size=10, indent=0)
            y -= 6

    c.save()
    buf.seek(0)
    return buf.read()


def _render_text_to_docx_bytes(title: str, text: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        raise RuntimeError('python-docx is required to generate DOCX: pip install python-docx')

    # Convert markdown to HTML and parse
    html = _markdown.markdown(text, extensions=["fenced_code", "tables"])
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')

    doc = Document()
    doc.add_heading(title, level=1)

    def add_paragraph_from_tag(tag):
        if tag.name in ('p', None):
            p = doc.add_paragraph()
            for item in tag.children:
                if getattr(item, 'name', None) == 'code':
                    run = p.add_run(item.get_text())
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                else:
                    run = p.add_run(item.get_text())
            return
        if tag.name == 'pre':
            p = doc.add_paragraph()
            run = p.add_run(tag.get_text())
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            return
        if tag.name in ('h1','h2','h3','h4'):
            level = 1 if tag.name=='h1' else 2 if tag.name=='h2' else 3
            doc.add_heading(tag.get_text(), level=level)
            return
        if tag.name == 'ul':
            for li in tag.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text(), style='List Bullet')
            return
        if tag.name == 'ol':
            for li in tag.find_all('li', recursive=False):
                p = doc.add_paragraph(li.get_text(), style='List Number')
            return

    for el in soup.contents:
        # handle text node
        if getattr(el, 'name', None) is None:
            txt = el.strip()
            if txt:
                doc.add_paragraph(txt)
        else:
            add_paragraph_from_tag(el)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


@app.route('/download/all/<fmt>')
def download_all(fmt):
    # Read blogs from server-side file if available
    blogs = []
    fpath = session.get('last_blogs_file')
    if fpath and os.path.exists(fpath):
        try:
            with open(fpath, 'r', encoding='utf-8') as fh:
                blogs = json.load(fh)
        except Exception:
            blogs = []
    else:
        try:
            blogs = json.loads(session.get('last_blogs', '[]'))
        except Exception:
            blogs = []
    filename = f"blogs.{fmt}"
    if fmt == 'json':
        content = json.dumps(blogs, ensure_ascii=False, indent=2)
        mimetype = 'application/json'
    elif fmt == 'csv':
        out_lines = []
        header = ['title', 'details', 'body', 'date']
        out_lines.append(','.join(header))
        for b in blogs:
            def esc(s):
                return '"' + (s or '').replace('"', '""') + '"'
            out_lines.append(','.join([esc(b.get('title')), esc(b.get('details')), esc(b.get('body')), esc(b.get('date'))]))
        content = '\n'.join(out_lines)
        mimetype = 'text/csv'
    else:
        # For other formats, concatenate files separated by a divider
        parts = []
        for i, b in enumerate(blogs, start=1):
            title = b.get('title','')
            if fmt in ('md','markdown'):
                parts.append(f"# {title}\n\n{b.get('body','')}\n\n---\n")
            elif fmt == 'html':
                parts.append(f"<h1>{title}</h1>\n{b.get('body_html','')}<hr/>")
            else:
                parts.append(f"{title}\n\n{b.get('body','')}\n\n---\n")
        content = '\n'.join(parts)
        mimetype = 'text/plain' if fmt in ('txt','md') else 'text/html'

    resp = make_response(content)
    resp.headers.set('Content-Type', mimetype + '; charset=utf-8')
    resp.headers.set('Content-Disposition', f'attachment; filename="{filename}"')
    return resp


@app.route("/api/chat", methods=["POST"])
def api_chat():
    """Enhanced Chat API endpoint for AI-powered article editing and conversational responses."""
    try:
        data = request.get_json()
        if not data or 'message' not in data:
            return json.dumps({"error": "No message provided"}), 400

        user_message = data['message'].strip()
        if not user_message:
            return json.dumps({"error": "Message cannot be empty"}), 400

        # Allow an optional per-request API key for chat.
        api_key = (CHAT_ASSISTANT_API_KEY or GEMINI_API_KEY)
        if api_key:
            api_key = api_key.strip()

        # Generate a concise assistant reply using the dedicated chat helper.
        try:
            response_text = llm_chat(user_message, api_key=api_key)
        except Exception as e:
            return json.dumps({"error": f"LLM chat error: {e}"}), 500

        # Return the model reply (trim to reasonable length)
        reply = response_text.strip() if response_text else "Unable to generate a response"
        return json.dumps({"response": reply}), 200
    except Exception as e:
        print(f"Chat API Error: {e}")
        return json.dumps({"error": str(e)}), 500

@app.route("/api/save-blog-content", methods=["POST"])
def save_blog_content():
    """Save updated blog content to session and database if user is logged in."""
    try:
        data = request.get_json()
        if not data or 'idx' not in data or 'body_html' not in data:
            return json.dumps({"error": "Missing idx or body_html"}), 400

        blog_idx = data['idx']
        updated_body_html = data['body_html']
        
        # Update in session (for non-authenticated users)
        try:
            blogs = json.loads(session.get('last_blogs', '[]'))
            if 0 <= blog_idx < len(blogs):
                blogs[blog_idx]['body_html'] = updated_body_html
                # Also update body as plain text
                blogs[blog_idx]['body'] = updated_body_html
                session['last_blogs'] = json.dumps(blogs)
                session.modified = True
        except Exception as e:
            print(f"Session update error: {e}")
        
        # Update in database if user is authenticated
        if current_user.is_authenticated:
            try:
                user_blogs = current_user.blogs.all()
                if 0 <= blog_idx < len(user_blogs):
                    blog = user_blogs[blog_idx]
                    blog.body_html = updated_body_html
                    blog.body = updated_body_html
                    db.session.commit()
            except Exception as e:
                print(f"Database update error: {e}")
                db.session.rollback()
        
        return json.dumps({"success": True, "message": "Blog content saved"}), 200
    except Exception as e:
        print(f"Save Blog Content Error: {e}")
        return json.dumps({"error": str(e)}), 500

@app.route("/api/user-stats", methods=["GET"])
def get_user_stats():
    """Get user statistics from database."""
    if not current_user or not getattr(current_user, 'is_authenticated', False):
        return json.dumps({'total_blogs': 0, 'total_downloads': 0}), 200, {'Content-Type': 'application/json'}
    
    try:
        # Simply count blogs from the database
        blog_count = UserBlog.query.filter_by(user_id=current_user.id).count()
        return json.dumps({
            'total_blogs': blog_count,
            'total_downloads': 0
        }), 200, {'Content-Type': 'application/json'}
    except Exception as e:
        print(f"Error fetching stats: {e}")
        return json.dumps({'total_blogs': 0, 'total_downloads': 0}), 200, {'Content-Type': 'application/json'}

@app.route("/api/trending-topics", methods=["GET"])
def get_trending_topics():
    """Get random trending topics from the database."""
    try:
        import random
        
        trending_topics = {
            'Tech': ['AI and Machine Learning Trends 2025', 'Web Development Best Practices', 'Cloud Computing Security', 'Cybersecurity for Startups', 'Python vs JavaScript Comparison', 'DevOps Pipeline Optimization', 'Blockchain Technology Explained', 'Quantum Computing Basics', 'IoT and Edge Computing', 'API Design Patterns'],
            'Beauty': ['Skincare Routine for Beginners', 'Natural Beauty Products Guide', 'Makeup Tips for Sensitive Skin', 'Anti-Aging Secrets from Experts', 'Hair Care Tips for Healthy Hair', 'Cruelty-Free Cosmetics Review', 'Wellness and Beauty Connection', 'K-Beauty Trends 2025', 'Sustainable Beauty Products', 'DIY Face Masks'],
            'Education': ['Online Learning Effectiveness', 'Study Techniques That Work', 'Educational Apps Review', 'Teaching Methods for Digital Age', 'Student Productivity Hacks', 'Online Courses Worth Taking', 'Critical Thinking Development', 'STEM Education Trends', 'Language Learning Tips', 'Educational Technology'],
            'Gaming': ['Top Gaming Trends 2025', 'Best PC Games of the Year', 'Gaming Setups Guide', 'Esports Career Opportunities', 'Indie Game Reviews', 'Mobile Gaming Tips', 'VR Gaming Experience', 'Gaming Streaming Guide', 'Game Development Basics', 'Gaming Console Comparison'],
            'Health': ['Fitness Goals Achievement Guide', 'Mental Health Awareness', 'Nutrition and Diet Tips', 'Workout Routines at Home', 'Sleep Quality Improvement', 'Stress Management Techniques', 'Preventive Health Measures', 'Yoga Benefits', 'Meditation Guide', 'Holistic Health Approach'],
            'Travel': ['Budget Travel Tips', 'Hidden Gems Destinations', 'Travel Packing Guide', 'Best Time to Travel', 'Travel Safety Essentials', 'Adventure Travel Planning', 'Cultural Experience Guide', 'Solo Travel Stories', 'Travel Photography Tips', 'Digital Nomad Guide'],
            'Lifestyle': ['Minimalist Living Guide', 'Work-Life Balance Tips', 'Daily Habits for Success', 'Personal Growth Strategies', 'Home Organization Ideas', 'Sustainable Living Tips', 'Time Management Mastery', 'Productivity Hacks', 'Morning Routine Benefits', 'Digital Detox Guide'],
            'Business': ['Startup Business Ideas', 'Entrepreneurship Guide', 'Marketing Strategies 2025', 'Leadership Skills Development', 'Financial Management Tips', 'Business Automation', 'Remote Work Best Practices', 'Scaling a Business', 'Customer Retention', 'Business Analytics']
        }
        
        # Get a random selection from each category
        result = {}
        for category, topics in trending_topics.items():
            result[category] = random.sample(topics, min(3, len(topics)))
        
        return json.dumps(result), 200
    except Exception as e:
        print(f"Trending Topics Error: {e}")
        return json.dumps({"error": str(e)}), 500

@app.route("/_debug_config")
def _debug_config():
    """Temporary debug endpoint to inspect LLM configuration."""
    cfg = {
        "MODEL": MODEL,
        "GEMINI_API_KEY_present": bool(GEMINI_API_KEY),
    }
    # Mask keys for safety
    return json.dumps(cfg, indent=2)

# ===== AUTHENTICATION ROUTES =====

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "").strip()
        remember = request.form.get("remember")

        if not email or not password:
            flash("Email and password are required", "error")
            return redirect(url_for("login"))

        user = User.query.filter_by(email=email).first()
        if user and user.check_password(password):
            login_user(user, remember=bool(remember))
            flash(f"Welcome back, {user.username}!", "success")
            return redirect(url_for("index"))
        else:
            flash("Invalid email or password", "error")

    return render_template("login.html")

@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        first_name = request.form.get("first_name", "").strip()
        last_name = request.form.get("last_name", "").strip()
        username = request.form.get("username", "").strip()
        email = request.form.get("email", "").strip()
        password = request.form.get("password", "").strip()
        terms = request.form.get("terms")

        # Validation
        if not all([first_name, username, email, password]):
            flash("All fields are required", "error")
            return redirect(url_for("register"))

        if not terms:
            flash("You must agree to the terms and conditions", "error")
            return redirect(url_for("register"))

        # Check if user already exists
        if User.query.filter_by(username=username).first():
            flash("Username already exists", "error")
            return redirect(url_for("register"))

        if User.query.filter_by(email=email).first():
            flash("Email already registered", "error")
            return redirect(url_for("register"))

        # Create new user
        user = User(username=username, email=email)
        user.set_password(password)
        
        try:
            db.session.add(user)
            db.session.commit()
            flash("Account created successfully! Please log in.", "success")
            return redirect(url_for("login"))
        except Exception as e:
            db.session.rollback()
            flash(f"An error occurred during registration: {str(e)}", "error")
            return redirect(url_for("register"))

    return render_template("register.html")

@app.route("/logout")
@login_required
def logout():
    logout_user()
    flash("You have been logged out", "success")
    return redirect(url_for("index"))

if __name__ == "__main__":
    app.run(debug=True)
